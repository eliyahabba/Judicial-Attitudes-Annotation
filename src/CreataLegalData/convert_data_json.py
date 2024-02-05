import json
import os
import re
import sys

import Levenshtein
import docx
import pandas as pd
from tqdm import tqdm
from pathlib import Path

from src.LegalAnnotationSentencesPage.FileText import FileText
from src.LegalAnnotationSentencesPage.TextDisplayer import TextDisplayer

SRC = Path(__file__).parents[1] / 'src'
sys.path.append(str(SRC))


VERDICT_PATH = r'C:\Users\t-eliyahabba\OneDrive - Microsoft\Documents\לימודים\מעבדה\DATA\base\data_by_type\sentencing_decisions'
VERDICT_PATH2 = r'C:\Users\t-eliyahabba\OneDrive - Microsoft\Documents\לימודים\מעבדה\DATA\base\origin_data\only_docx_files'


def map_sentences_to_paragraphs(paragraphs):
    sentences_tp_pars = {}
    global_sentences_index = 0
    for i, paragraph in enumerate(paragraphs):
        sentences = [sentence.strip() for sentence in split_paragraph_to_sentence(paragraph) if sentence]
        for j, sentence in enumerate(sentences):
            sentences_tp_pars[global_sentences_index + j] = i
        global_sentences_index += len(sentences)
    return sentences_tp_pars


def concat_sentences(lst):
    i = 0
    merged_lst = []
    special_chars = ["!", "?", "."]
    while i < len(lst):
        # Check if the next element consists only of special characters
        if i < len(lst) - 1 and all(char in special_chars for char in lst[i + 1]):
            # Merge the current element and the next element
            merged_element = lst[i] + lst[i + 1]
            # Add the merged element to the new list
            merged_lst.append(merged_element)
            # Skip the next element
            i += 1
        else:
            # Add the current element to the new list
            merged_lst.append(lst[i])
        i += 1
    return merged_lst


def split_paragraph_to_sentence(paragraph):
    split_chars = ["!", "?", "."]

    # Use a regular expression to match any of the split characters
    split_regex = r"(?<!\d)(" + "|".join(re.escape(char) for char in split_chars) + ")"
    result = re.split(split_regex, paragraph)
    # take only the sentences that are not empty anf have characters in them (not only special chars)
    result = [sentence for sentence in result if sentence and any(char.isalpha() for char in sentence)]
    # concat sentences that are split by mistake

    # paragraph_result = concat_sentences(result)

    sentences = [sentence.strip() for sentence in result if sentence]
    return sentences


def create_json_from_docx(file_name, sentence_id, main_paragraph, main_sentence=None):
    file_text = FileText(file_name, sentence_id)
    file_name = file_text.file_name + ".docx"

    file_path = os.path.join(VERDICT_PATH, file_name)
    if not os.path.exists(file_path):
        file_path = os.path.join(VERDICT_PATH2, file_name)
        if not os.path.exists(file_path):
            return -1, file_path

    doc = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    paragraphs = [par for par in paragraphs if par]
    sentences = [split_paragraph_to_sentence(paragraph) for paragraph
                 in paragraphs]

    text_displayer = TextDisplayer(file_text)

    par_ind = text_displayer.find_relevant_paragraph(paragraphs, main_paragraph)
    sen_ind_in_par = text_displayer.find_relevant_paragraph(
        [sentence.strip() for sentence in split_paragraph_to_sentence(paragraphs[par_ind]) if sentence]
        , main_sentence)

    global_sentence_index = 0
    for i in range(par_ind):
        global_sentence_index += len(sentences[i])
    global_sentence_index += sen_ind_in_par

    # save flatten sentences
    flatten_sentences = [sentence for par in sentences for sentence in par]
    dis = Levenshtein.distance(flatten_sentences[global_sentence_index].strip(), main_sentence.strip()) / len(
        main_sentence)
    if flatten_sentences[global_sentence_index] != main_sentence:
        dis = Levenshtein.distance(flatten_sentences[global_sentence_index].strip(),
                                   main_sentence.strip()) / len(main_sentence)
        # trunf dis to 2 digits
        dis = int(dis * 100) / 100

    sentences_tp_pars = map_sentences_to_paragraphs(paragraphs)

    pars_with_sentence = {}
    for sen_ind, sentence in enumerate(sentences):
        # take all the sentences with same paragraph index (from sentences_to_paragraphs),
        # and put them in a list with the same order as in the original text
        if sentences_tp_pars[sen_ind] in pars_with_sentence:
            pars_with_sentence[sentences_tp_pars[sen_ind]].append(sentence)
        else:
            pars_with_sentence[sentences_tp_pars[sen_ind]] = [sentence]
    sentences_tp_sen = {}
    q = 0
    for i, paragraph in enumerate(sentences):
        sentences_in_par = paragraph
        for j, sentence in enumerate(sentences_in_par):
            # sentences_tp_pars[global_sentences_index + j] = i
            # add i, j to sentence name
            sentence_k = sentence + "_" + str(i) + "_" + str(j)

            sentences_tp_sen[sentence_k] = q
            q += 1
        # global_sentences_index += len(sentences_in_par)

    file_name = file_text.file_name

    path = Path(__file__).parents[2] / "Data" / "json_docx_v2"
    os.makedirs(path, exist_ok=True)
    if not os.path.exists(os.path.join(path, file_name + f".json")):
        with open(os.path.join(path, file_name + f".json"), "w", encoding="utf-8") as f:
            json.dump({"title": file_name,
                       "paragraphs": [{"paragraph_index": i,
                                       "sentences": [{"sentence": sentence,
                                                      "sentence_id": str(j)
                                                         , "paragraph_index": i,
                                                      "global_sentence_index": sentences_tp_sen[
                                                          sentence + "_" + str(i) + "_" + str(j)]}
                                                     for j, sentence in enumerate(paragraph)]}
                                      for i, paragraph in enumerate(sentences)],
                       "main_sentence_from_spike": [main_sentence],
                       "main_sentence_from_docx": [flatten_sentences[global_sentence_index]],
                       "main_sentence_id_in_docx": [str(par_ind) + "_" + str(sen_ind_in_par)],
                       "global_sentence_index_docx": [global_sentence_index],
                       "global_sentence_index_spike": [sentence_id],
                       "score_of_similarity": [dis]},
                      f, ensure_ascii=False, indent=4)
    else:
        # add to existing file
        data = json.load(open(os.path.join(path, file_name + f".json"), "r", encoding="utf-8"))
        if global_sentence_index in data["global_sentence_index_docx"] or sentence_id in data[
            "global_sentence_index_spike"]:
            return dis, None

        data["main_sentence_from_spike"].append(main_sentence)
        data["main_sentence_from_docx"].append(flatten_sentences[global_sentence_index])
        data["main_sentence_id_in_docx"].append(str(par_ind) + "_" + str(sen_ind_in_par))
        data["global_sentence_index_docx"].append(global_sentence_index)
        data["global_sentence_index_spike"].append(sentence_id)
        data["score_of_similarity"].append(dis)
        # write the new data to the file
        with open(os.path.join(path, file_name + f".json"), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)

    return dis, None


def write_file_without_main_sentence(path, file_name):
    file_path = os.path.join(path, file_name + ".docx")
    doc = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    paragraphs = [par for par in paragraphs if par]
    sentences = [split_paragraph_to_sentence(paragraph) for paragraph
                 in paragraphs]

    sentences_tp_sen = {}
    q = 0
    for i, paragraph in enumerate(sentences):
        sentences_in_par = paragraph
        for j, sentence in enumerate(sentences_in_par):
            # sentences_tp_pars[global_sentences_index + j] = i
            # add i, j to sentence name
            sentence_k = sentence + "_" + str(i) + "_" + str(j)

            sentences_tp_sen[sentence_k] = q
            q += 1
        # global_sentences_index += len(sentences_in_par)

    path = Path(__file__).parents[2] / "Data" / "json_docx_v2"
    os.makedirs(path, exist_ok=True)
    if not os.path.exists(os.path.join(path, file_name + f".json")):
        with open(os.path.join(path, file_name + f".json"), "w", encoding="utf-8") as f:
            json.dump({"title": file_name,
                       "paragraphs": [{"paragraph_index": i,
                                       "sentences": [{"sentence": sentence,
                                                      "sentence_id": str(j)
                                                         , "paragraph_index": i,
                                                      "global_sentence_index": sentences_tp_sen[
                                                          sentence + "_" + str(i) + "_" + str(j)]}
                                                     for j, sentence in enumerate(paragraph)]}
                                      for i, paragraph in enumerate(sentences)],
                       "main_sentence_from_spike": [],
                       "main_sentence_from_docx": [],
                       "main_sentence_id_in_docx": [],
                       "global_sentence_index_docx": [],
                       "global_sentence_index_spike": [],
                       "score_of_similarity": []},
                      f, ensure_ascii=False, indent=4)


def read_csv(df):
    problem_files = []
    c = 0
    diss = []
    files_in_spike_data = df['title'].unique()

    for i, row in tqdm(df.iterrows(), total=len(df)):
        try:
            main_sentence = row['origin_sentence']
            file_name = row['title']
            # if file_name != 'k00018098p':
            #     continue
            main_paragraph = row['paragraph_text']
            sentence_id = row['sentence_id']
            dis, file_name = create_json_from_docx(file_name, sentence_id, main_paragraph, main_sentence)
            if dis < 0:
                problem_files.append(file_name)
        except Exception as e:
            c += 1
            print(e, "\n", file_name)
            continue

    print("There are {} problem files out of {}".format(len(problem_files), len(df)))
    print("There art {} sentences out of {}".format(c, len(df)))
    # count how many values in diss >0.2
    print("There are {} sentences out of {} with distance >0.2".format(len([dis for dis in diss if dis > 0.2]),
                                                                       len(diss)))
    # print(max(diss))

    verdicts = os.listdir(VERDICT_PATH)
    verdicts = [verdict.split(".")[0] for verdict in verdicts]
    verdicts = [verdict for verdict in verdicts if verdict not in files_in_spike_data]

    for file in tqdm(verdicts):
        write_file_without_main_sentence(VERDICT_PATH, file)


if __name__ == '__main__':
    from pathlib import Path
    import streamlit as st
    st.session_state['files_number']=0
    df_path = Path(__file__).parents[2] / "Data" / "spike_results_v8.csv"
    df = pd.read_csv(df_path)
    read_csv(df)
