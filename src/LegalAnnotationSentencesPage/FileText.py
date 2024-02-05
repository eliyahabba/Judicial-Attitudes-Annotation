import json
import os
import warnings

import docx
import streamlit as st

warnings.filterwarnings('ignore')


class FileText:
    def __init__(self, file_name, sen_id):
        self.data = None
        self.sentences = None
        self.file_name = file_name
        self.sen_id = sen_id
        self.file_path = None
        self.paragraphs = None

    def find_file_path(self):
        file_name = self.file_name + ".json"
        st.session_state["file_name"] = file_name
        st.session_state["sentence_index"] = self.sen_id
        self.file_path = st.session_state.jsons_path / file_name
        assert os.path.exists(self.file_path), f'file {self.file_path} does not exist'

    def get_paragraphs_old(self):
        doc = docx.Document(self.file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        paragraphs = [par for par in paragraphs if par]
        sentences = [sentence.strip() + ("." if not sentence.endswith((":", "-", "!", "?")) else "") for paragraph in
                     paragraphs for sentence in paragraph.split(".") if
                     sentence]
        self.paragraphs = paragraphs
        self.sentences = sentences

    def get_paragraphs(self):
        with open(self.file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        paragraphs = [par for par in data["paragraphs"]]
        sentences = sum([len(par["sentences"]) for par in paragraphs])
        self.paragraphs = paragraphs
        self.sentences = sentences

        self.data = data
